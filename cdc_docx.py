"""Cahier des charges : export Word (.docx) a la charte NSN Industrie, et relecture.

Le gabarit applique le design system du groupe (c:/projets/ID visuel,
DESIGN-nsn.md + maquette-rapport.html), pas une charte inventee :

- encre #0F2329 (le noir cyane du logotype), corps #33555F, libelles #647E86,
  filets #CFD8DB / #AFBCC1, fonds enfonces #E5EBEC / #F3F6F7 ;
- titrage en Bahnschrift SemiCondensed (le DIN de Windows, deja installe),
  gras, capitales, lettres espacees ; corps en Segoe UI ; donnees en Consolas ;
- la couleur ne dit qu'une chose : la societe. La barre tricolore sous le
  bandeau (bleu Usinage, vert Sopranzi, jaune R.A.C) est la signature du
  groupe, reprise de la maquette de rapport ;
- angles a 0, aucune ombre, filets visibles plutot que boites flottantes.

Les CAPITALES d'affichage passent par la propriete Word w:caps, JAMAIS par une
transformation du texte : les caracteres stockes restent intacts, c'est ce qui
garantit l'aller-retour fidele.

Aller-retour :
    build(chantier)  -> (octets, nom de fichier)
    parse(octets)    -> {"chantier_id": str|None, "sections": [{id,titre,corps}]}

Le rattachement des sections au retour se fait dans cet ordre : identifiants
ranges dans docProps/custom.xml (invisibles dans Word, conserves par un
enregistrement normal) ; a defaut le titre normalise ; a defaut la position.
Une section ajoutee dans Word arrive sans identifiant, l'appli lui en cree un.
Une section supprimee dans Word disparait de l'appli. En suivi de modifications
Word, on lit l'etat FINAL (le texte supprime vit dans <w:delText>, pas relu).
"""

from __future__ import annotations

import io
import re
import unicodedata
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

# ---- jetons du design system (ID visuel / tokens) -------------------------- #
INK = RGBColor(0x0F, 0x23, 0x29)          # encre du logotype
BODY = RGBColor(0x33, 0x55, 0x5F)         # texte courant
STRONG = RGBColor(0x08, 0x14, 0x18)       # texte appuye
MUTED = RGBColor(0x64, 0x7E, 0x86)        # libelles
MUTED_SOFT = RGBColor(0x87, 0x98, 0xA0)   # libelles discrets
ON_DARK_SOFT = RGBColor(0x9F, 0xB0, 0xB5)  # texte secondaire sur encre
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_INK = "0F2329"
HEX_HAIR = "CFD8DB"
HEX_HAIR2 = "AFBCC1"
HEX_SUNK = "E5EBEC"
HEX_SOFT = "F3F6F7"
HEX_USI = "0D5FA6"                        # NSN Usinage
HEX_SOP = "0A8A3F"                        # NSN Sopranzi
HEX_RAC = "EDA400"                        # NSN R.A.C

DIN = "Bahnschrift SemiCondensed"         # repli natif du D-DIN Condensed
SANS = "Segoe UI"
MONO = "Consolas"

SOCIETE = {"nom": "NSN", "adresse": "972 Avenue du 19 Mars 1962 · 38540 Heyrieux",
           "adresse_lignes": ["972 Avenue du 19 Mars 1962", "38540 Heyrieux"]}
HEX_BORD = "333333"                       # bordures du cartouche, comme le RFF
HEX_NOIR = "1A1A1A"                       # en-tetes de tableaux, texte blanc
PROPRIETE = ("Ce document est la propriété de NSN Industrie. Il ne peut être "
             "reproduit ni communiqué à un tiers sans autorisation écrite.")

# La marque officielle (geometrie du site, relevee dans ID visuel), rendue en
# PNG a fond transparent : negatif pour le bandeau encre, encre pour la garde.
_ASSETS = Path(__file__).resolve().parent / "assets"
MARK_NEGATIF = _ASSETS / "nsn-mark-negatif.png"
MARK_ENCRE = _ASSETS / "nsn-mark.png"

STYLE_TITRE = "Heading1"
PROP_CHANTIER = "SuiviChantierId"
PROP_SECTIONS = "SuiviSectionIds"

CDC_STATUT_LBL = {"brouillon": "Brouillon", "en_validation": "En validation",
                  "valide": "Validé", "obsolete": "Obsolète"}
CDC_STATUT_EN = {"brouillon": "Draft", "en_validation": "Under review",
                 "valide": "Approved", "obsolete": "Obsolete"}


# --------------------------------------------------------------------------- #
# Briques bas niveau
# --------------------------------------------------------------------------- #
def _run(p, text, *, bold=False, italic=False, size=9.5, color=BODY, font=SANS,
         caps=False, ls=0):
    """Un run type. caps = capitales d'AFFICHAGE (w:caps), le texte ne change
    pas ; ls = interlettrage en vingtiemes de point (w:spacing)."""
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    r.font.name = font
    r.font.color.rgb = color
    rpr = r._r.get_or_add_rPr()
    if caps:
        rpr.append(OxmlElement("w:caps"))
    if ls:
        el = OxmlElement("w:spacing")
        el.set(qn("w:val"), str(ls))
        rpr.append(el)
    return r


def _shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _shade_p(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)


def _p_border(p, *, edge="bottom", sz=4, color=HEX_HAIR, space=3):
    pPr = p._p.get_or_add_pPr()
    bdr = pPr.find(qn("w:pBdr"))
    if bdr is None:
        bdr = OxmlElement("w:pBdr")
        pPr.append(bdr)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:color"), color)
    e.set(qn("w:space"), str(space))
    bdr.append(e)


def _tbl_borders(table, spec):
    """spec = {edge: (sz, couleur) | None}. Les aretes absentes restent nulles :
    la charte trace des filets horizontaux, pas des grilles."""
    tblPr = table._element.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        if spec.get(edge):
            sz, color = spec[edge]
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(sz))
            e.set(qn("w:color"), color)
        else:
            e.set(qn("w:val"), "none")
        b.append(e)
    tblPr.append(b)


def _row_h(row, mm, exact=True):
    row.height = Mm(mm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if exact else WD_ROW_HEIGHT_RULE.AT_LEAST


def _grid(table, mms):
    """Verrouille les largeurs de colonnes. Poser tcW ne suffit pas : en
    disposition fixe Word arbitre sur w:tblGrid, et la grille par defaut est
    equirepartie — c'est elle qui tronquait la case INDICE/PAGE du bandeau."""
    tbl = table._element
    for gc, mm in zip(tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol")), mms):
        gc.set(qn("w:w"), str(int(mm * 56.693)))
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(sum(mms) * 56.693)))
    for row in table.rows:
        for cell, mm in zip(row.cells, mms):
            cell.width = Mm(mm)


def _clear(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def _page_field(paragraph, *, size=7.5, color=ON_DARK_SOFT):
    """« Page X / Y » en champs Word, stylable."""
    def fld(instr):
        f = OxmlElement("w:fldSimple")
        f.set(qn("w:instr"), instr)
        rr = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        c = OxmlElement("w:color")
        c.set(qn("w:val"), f"{color:06X}" if isinstance(color, int) else str(color))
        rpr.append(c)
        szel = OxmlElement("w:sz")
        szel.set(qn("w:val"), str(int(size * 2)))
        rpr.append(szel)
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), DIN)
        rf.set(qn("w:hAnsi"), DIN)
        rpr.append(rf)
        rr.append(rpr)
        t = OxmlElement("w:t")
        t.text = "1"
        rr.append(t)
        f.append(rr)
        return f
    lbl = RGBColor.from_string(str(color)) if not isinstance(color, RGBColor) else color
    _run(paragraph, "Page ", size=size, color=lbl)
    paragraph._p.append(fld("PAGE"))
    _run(paragraph, " / ", size=size, color=lbl)
    paragraph._p.append(fld("NUMPAGES"))


def _toc(doc):
    p = doc.add_paragraph()
    p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-1" \h \z \u')
    holder = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Clic droit → « Mettre à jour les champs » pour générer la table."
    holder.append(t)
    fld.append(holder)
    p._p.append(fld)


# --------------------------------------------------------------------------- #
# En-tete et pied de page (repetes sur chaque page)
# --------------------------------------------------------------------------- #
def _entete(section, ch, cdc):
    """Cartouche d'en-tete du gabarit RFF : bordures noires, marque + adresse,
    titre centre, cases REV. et PAGE, puis le bandeau bilingue CHANTIER /
    REDACTEUR. La barre tricolore du groupe reste dessous."""
    hdr = section.header
    hdr.is_linked_to_previous = False
    for p in list(hdr.paragraphs):
        p._element.getparent().remove(p._element)

    GRILLE = {e: (8, HEX_BORD) for e in
              ("top", "left", "bottom", "right", "insideH", "insideV")}
    cart = hdr.add_table(rows=2, cols=5, width=Mm(186))
    cart.autofit = False
    _tbl_borders(cart, GRILLE)
    _grid(cart, [26, 34, 97, 11, 18])
    for row in cart.rows:
        _row_h(row, 7.5, exact=False)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # marque officielle (encre, repere menthe) sur fond blanc
    lc = cart.cell(0, 0).merge(cart.cell(1, 0))
    p = _clear(lc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if MARK_ENCRE.exists():
        try:
            p.add_run().add_picture(str(MARK_ENCRE), height=Mm(10))
        except Exception:
            _run(p, "NSN", bold=True, size=12, color=INK, font=DIN)

    # nom + adresse
    ac = cart.cell(0, 1).merge(cart.cell(1, 1))
    p = _clear(ac)
    _run(p, SOCIETE["nom"], bold=True, size=9, color=INK)
    for ligne in SOCIETE["adresse_lignes"]:
        pp = ac.add_paragraph()
        pp.paragraph_format.space_before = Pt(0)
        pp.paragraph_format.space_after = Pt(0)
        _run(pp, ligne, size=7, color=BODY)

    # titre + N°
    tc = cart.cell(0, 2).merge(cart.cell(1, 2))
    p = _clear(tc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Cahier des Charges", bold=True, size=12, color=INK)
    p2 = tc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(0)
    _run(p2, "N° CDC : ", size=8, color=MUTED)
    _run(p2, cdc.get("reference") or ch.get("id", ""), bold=True, size=9, color=INK)

    # REV. / PAGE : libelle en haut, valeur en bas, comme le RFF
    _cellule(cart.cell(0, 3), "REV.", bold=True, size=6.5, color=MUTED, centre=True)
    _cellule(cart.cell(1, 3), cdc.get("indice") or "A", bold=True, size=11,
             color=INK, centre=True)
    _cellule(cart.cell(0, 4), "PAGE", bold=True, size=6.5, color=MUTED, centre=True)
    pc = cart.cell(1, 4)
    p = _clear(pc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _page_field(p, size=7, color="647E86")

    # bandeau bilingue CHANTIER / REDACTEUR
    band = hdr.add_table(rows=1, cols=2, width=Mm(186))
    band.autofit = False
    _tbl_borders(band, GRILLE)
    _grid(band, [93, 93])
    _row_h(band.rows[0], 5.5, exact=False)
    for cell, (fr, en, val) in zip(band.rows[0].cells, [
        ("CHANTIER / ", "PROJECT", " : " + (ch.get("titre") or "")),
        ("RÉDACTEUR / ", "AUTHOR", " : " + (cdc.get("redacteur") or "")),
    ]):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = _clear(cell)
        _run(p, fr, bold=True, size=7, color=INK)
        _run(p, en, bold=True, italic=True, size=7, color=MUTED)
        _run(p, val, size=7, color=BODY)

    tail = hdr.add_paragraph()
    tail.paragraph_format.space_before = Pt(0)
    tail.paragraph_format.space_after = Pt(0)
    _run(tail, "", size=4)


def _cellule(cell, texte, *, bold=False, size=8, color=BODY, centre=False):
    p = _clear(cell)
    if centre:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, texte, bold=bold, size=size, color=color)


def _pied(section, ch, cdc):
    ftr = section.footer
    ftr.is_linked_to_previous = False
    for p in list(ftr.paragraphs):
        p._element.getparent().remove(p._element)
    p = ftr.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    _p_border(p, edge="top", sz=4, color=HEX_HAIR, space=4)
    _run(p, SOCIETE["nom"], bold=True, size=6.5, color=MUTED, font=DIN, ls=24)
    _run(p, "  ·  " + SOCIETE["adresse"] + "  ·  ", size=6.5, color=MUTED_SOFT)
    _run(p, (cdc.get("reference") or "") + " — " + (ch.get("titre") or ""),
         size=6.5, color=MUTED_SOFT, italic=True)
    p2 = ftr.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    _run(p2, PROPRIETE, italic=True, size=6, color=MUTED_SOFT)


# --------------------------------------------------------------------------- #
# Blocs de garde
# --------------------------------------------------------------------------- #
def _eyebrow(doc, fr, en):
    """Intitule de bloc du gabarit RFF : FR gras noir, EN italique gris,
    filet sombre sous toute la largeur."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    _run(p, fr, bold=True, size=10, color=STRONG, font=DIN, caps=True, ls=24)
    _run(p, "  /  ", size=8, color=MUTED_SOFT, font=DIN)
    _run(p, en, italic=True, size=8, color=MUTED_SOFT)
    _p_border(p, sz=6, color="444444", space=3)
    return p


def _table_charte(doc, cols, rows, empty_rows=0):
    """cols = [(titre, largeur_mm, genre)], genre dans {label, body, mono, en,
    center}. Gabarit RFF : en-tete NOIR texte blanc, grille complete."""
    n = len(rows) + empty_rows + 1
    t = doc.add_table(rows=n, cols=len(cols))
    t.autofit = False
    _tbl_borders(t, {e: (8, HEX_BORD) for e in
                     ("top", "left", "bottom", "right", "insideH", "insideV")})
    _grid(t, [w for (_t, w, _g) in cols])
    body = list(rows) + [tuple("" for _ in cols) for _ in range(empty_rows)]

    _row_h(t.rows[0], 6.5, exact=False)
    for j, (titre, w, genre) in enumerate(cols):
        cell = t.cell(0, j)
        _shade_cell(cell, HEX_NOIR)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = _clear(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, titre, bold=True, size=8, color=WHITE, font=DIN, caps=True, ls=24)

    for ri, brow in zip(range(1, n), body):
        _row_h(t.rows[ri], 6, exact=False)
        for j, (_titre, _w, genre) in enumerate(cols):
            cell = t.cell(ri, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = _clear(cell)
            val = str(brow[j])
            if genre == "label":
                _run(p, val, bold=True, size=8, color=INK, font=DIN, caps=True, ls=20)
            elif genre == "mono":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run(p, val, size=8, color=BODY, font=MONO)
            elif genre == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run(p, val, bold=True, size=9, color=INK, font=DIN)
            elif genre == "en":
                _run(p, val, italic=True, size=7.5, color=MUTED_SOFT)
            else:
                _run(p, val, size=8.5, color=BODY)
    return t


# --------------------------------------------------------------------------- #
# Corps : mise en forme SANS toucher aux caracteres
# --------------------------------------------------------------------------- #
_RE_PUCE = re.compile(r"^\s*[-•*]\s+")
_RE_NUMEROTE = re.compile(r"^\s*\d+[.)]\s+")
_RE_COLONNES = re.compile(r"\S {3,}\S")
_RE_CONTINUATION = re.compile(r"^\s{6,}\S")
_RE_LABEL = re.compile(r"^([A-ZÀ-ÖØ-Þ0-9][A-ZÀ-ÖØ-Þ0-9 '’/()&.,’-]{2,70}?\s*:)(\s*)(.*)$")


def _indent(p, gauche_mm, retrait_mm=0.0, droite_mm=0.0):
    pf = p.paragraph_format
    pf.left_indent = Mm(gauche_mm)
    if retrait_mm:
        pf.first_line_indent = Mm(-retrait_mm)
    if droite_mm:
        pf.right_indent = Mm(droite_mm)


def _para_corps(doc, ligne: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2.5)
    pf.line_spacing = 1.12

    if not ligne.strip():                        # respiration
        _run(p, "", size=4)
        pf.space_after = Pt(0)
        return p

    if _RE_COLONNES.search(ligne) or _RE_CONTINUATION.match(ligne):
        # bloc aligne a la main -> panneau chasse fixe sur fond doux ;
        # les lignes consecutives forment visuellement un seul panneau.
        _run(p, ligne, size=8, color=BODY, font=MONO)
        _shade_p(p, HEX_SOFT)
        _indent(p, 3, droite_mm=3)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.25
        return p

    m_puce = _RE_PUCE.match(ligne) or _RE_NUMEROTE.match(ligne)
    if m_puce:
        marque = m_puce.group(0)
        _run(p, marque, bold=True, size=9.5, color=INK)
        _run(p, ligne[len(marque):], size=9.5, color=BODY)
        _indent(p, 8, 4)
        return p

    m = _RE_LABEL.match(ligne)
    if m and m.group(3).strip():                 # « LABEL : » suivi de texte
        _run(p, m.group(1), bold=True, size=9.5, color=INK)
        _run(p, m.group(2) + m.group(3), size=9.5, color=BODY)
        return p
    if m:                                        # « LABEL : » seul -> sous-titre
        pf.space_before = Pt(8)
        pf.space_after = Pt(3)
        _run(p, ligne, bold=True, size=9.5, color=INK, font=DIN, ls=24)
        return p

    _run(p, ligne, size=9.5, color=BODY)
    return p


# --------------------------------------------------------------------------- #
# Ecriture
# --------------------------------------------------------------------------- #
def _nom_fichier(ch, cdc) -> str:
    base = cdc.get("reference") or cdc.get("titre") or ch.get("titre") or "cahier_des_charges"
    base = unicodedata.normalize("NFKD", base).encode("ascii", "ignore").decode()
    base = re.sub(r"[^A-Za-z0-9._-]+", "_", base).strip("_") or "cahier_des_charges"
    return f"CDC_{base}_{cdc.get('indice', 'A')}.docx"


def _custom_props(chantier_id: str, section_ids: list) -> str:
    fmtid = "{D5CDD505-2E9C-101B-9397-08002B2CF9AE}"

    def prop(pid, nom, val):
        return (f'<property fmtid="{fmtid}" pid="{pid}" name="{escape(nom)}">'
                f"<vt:lpwstr>{escape(val)}</vt:lpwstr></property>")

    return ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"'
            ' xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
            + prop(2, PROP_CHANTIER, chantier_id or "")
            + prop(3, PROP_SECTIONS, ",".join(section_ids))
            + "</Properties>")


def _injecte_props(data: bytes, chantier_id: str, ids: list) -> bytes:
    """Ajoute docProps/custom.xml au paquet (python-docx ne l'expose pas)."""
    src = zipfile.ZipFile(io.BytesIO(data))
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for item in src.infolist():
            contenu = src.read(item.filename)
            if item.filename == "[Content_Types].xml":
                txt = contenu.decode("utf-8")
                if "custom-properties" not in txt:
                    txt = txt.replace("</Types>",
                                      '<Override PartName="/docProps/custom.xml" ContentType='
                                      '"application/vnd.openxmlformats-officedocument.'
                                      'custom-properties+xml"/></Types>')
                contenu = txt.encode("utf-8")
            elif item.filename == "_rels/.rels":
                txt = contenu.decode("utf-8")
                if "docProps/custom.xml" not in txt:
                    txt = txt.replace("</Relationships>",
                                      '<Relationship Id="rIdSuiviCustom" Target="docProps/custom.xml" '
                                      'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
                                      'relationships/custom-properties"/></Relationships>')
                contenu = txt.encode("utf-8")
            z.writestr(item, contenu)
        z.writestr("docProps/custom.xml", _custom_props(chantier_id, ids))
    return out.getvalue()


def build(ch: dict) -> tuple:
    """Rend (octets du .docx, nom de fichier) pour le cahier des charges."""
    cdc = ch.get("cdc")
    if not cdc:
        raise ValueError("Ce chantier n'a pas de cahier des charges.")
    sections = cdc.get("sections") or []

    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Mm(297), Mm(210)
    sec.top_margin, sec.bottom_margin = Mm(30), Mm(18)
    sec.left_margin = sec.right_margin = Mm(12)
    sec.header_distance = Mm(7)
    sec.footer_distance = Mm(9)
    normal = doc.styles["Normal"]
    normal.font.name = SANS
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = BODY

    _entete(sec, ch, cdc)
    _pied(sec, ch, cdc)

    # ---- page de garde ------------------------------------------------------
    # Le cartouche encadre du gabarit RFF : titre et N° centres dans une boite
    # a bordure noire epaisse, avec de l'air au-dessus et en dessous.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    _run(p, "", size=8)
    boite = doc.add_table(rows=1, cols=1)
    boite.autofit = False
    _tbl_borders(boite, {e: (16, "111111") for e in
                         ("top", "left", "bottom", "right")})
    _grid(boite, [150])
    boite.alignment = 1                       # WD_TABLE_ALIGNMENT.CENTER
    cell = boite.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = _clear(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    _run(p, "Cahier des Charges", bold=True, size=22, color=INK)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    _run(p2, "N° CDC : ", size=12, color=BODY)
    _run(p2, cdc.get("reference") or ch.get("id", ""), bold=True, size=13, color=INK)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(14)
    _run(p3, cdc.get("titre") or ch.get("titre") or "", size=10, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    _run(p, "", size=8)

    # ---- identification -----------------------------------------------------
    _eyebrow(doc, "Identification", "Document identification")
    statut = cdc.get("statut") or "brouillon"
    _table_charte(doc, [("Rubrique", 42, "label"), ("Valeur", 96, "body"),
                        ("Value", 48, "en")], [
        ("Référence", cdc.get("reference") or "", "Reference"),
        ("Indice", cdc.get("indice") or "A", "Revision"),
        ("Statut", CDC_STATUT_LBL.get(statut, statut), CDC_STATUT_EN.get(statut, "")),
        ("Rédacteur", cdc.get("redacteur") or "", "Author"),
        ("Créé le", str(cdc.get("date_creation") or ""), "Created"),
        ("Mis à jour le", str(cdc.get("date_maj") or ""), "Updated"),
        ("Validé par", cdc.get("valide_par") or "", str(cdc.get("date_validation") or "")),
    ])

    # ---- liste des revisions ------------------------------------------------
    _eyebrow(doc, "Liste des révisions", "Revision list")
    revs = [(r.get("indice") or "", str(r.get("date") or ""),
             r.get("auteur") or "", r.get("objet") or "")
            for r in cdc.get("revisions") or []]
    _table_charte(doc, [("Indice", 16, "center"), ("Date", 24, "mono"),
                        ("Auteur", 30, "body"), ("Objet de la révision", 116, "body")],
                  revs, empty_rows=max(0, 3 - len(revs)))

    # ---- approbation --------------------------------------------------------
    # Le cartouche a visas des documents industriels : redaction, verification,
    # approbation. Les cases restent a signer sur papier ou en PDF.
    _eyebrow(doc, "Approbation", "Approval")
    appro = doc.add_table(rows=4, cols=4)
    appro.autofit = False
    _tbl_borders(appro, {e: (8, HEX_BORD) for e in
                         ("top", "left", "bottom", "right", "insideH", "insideV")})
    _grid(appro, [24, 54, 54, 54])
    _row_h(appro.rows[0], 6.5, exact=False)
    for j, titre in enumerate(("", "Rédaction", "Vérification", "Approbation")):
        cell = appro.cell(0, j)
        _shade_cell(cell, HEX_NOIR)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = _clear(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, titre, bold=True, size=8, color=WHITE, font=DIN, caps=True, ls=24)
    valeurs = {
        1: (cdc.get("redacteur") or "", str(cdc.get("date_creation") or "")),
        2: ("", ""),
        3: (cdc.get("valide_par") or "", str(cdc.get("date_validation") or "")),
    }
    for ri, lbl in ((1, "Nom"), (2, "Date"), (3, "Visa")):
        _row_h(appro.rows[ri], 13 if lbl == "Visa" else 6, exact=(lbl == "Visa"))
        cell = appro.cell(ri, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _run(_clear(cell), lbl, bold=True, size=7.5, color=MUTED, font=DIN,
             caps=True, ls=24)
        for j in (1, 2, 3):
            cell = appro.cell(ri, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = _clear(cell)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if lbl == "Nom":
                _run(p, valeurs[j][0], size=8.5, color=BODY)
            elif lbl == "Date":
                _run(p, valeurs[j][1], size=8, color=BODY, font=MONO)

    # ---- parties prenantes --------------------------------------------------
    parties = cdc.get("parties_prenantes") or []
    if parties:
        _eyebrow(doc, "Parties prenantes", "Stakeholders")
        lignes = []
        for x in parties:
            if isinstance(x, dict):
                lignes.append((x.get("nom") or "", x.get("role") or "", x.get("societe") or ""))
            else:
                lignes.append((str(x), "", ""))
        _table_charte(doc, [("Nom", 60, "label"), ("Rôle", 66, "body"),
                            ("Société", 60, "body")], lignes)

    # ---- pilotage : planning et risques, GENERES depuis le chantier ---------
    # Ces deux blocs ne sont pas des sections editables : le plan et le
    # registre de risques vivent dans l'appli, les recopier dans le texte
    # creerait deux verites. Ils sont rendus comme l'identification (hors
    # style Titre 1), donc invisibles pour la relecture parse().
    doc.add_page_break()
    _eyebrow(doc, "Planning et jalons", "Schedule and milestones")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _run(p, "DÉBUT ", bold=True, size=7.5, color=MUTED, font=DIN, ls=24)
    _run(p, str(ch.get("date_debut") or "—"), size=8.5, color=BODY, font=MONO)
    _run(p, "    ÉCHÉANCE ", bold=True, size=7.5, color=MUTED, font=DIN, ls=24)
    _run(p, str(ch.get("echeance") or "—"), size=8.5, color=BODY, font=MONO)
    _run(p, "    ·    le planning de référence (chemin critique, jours ouvrés) "
            "vit dans l'appli de suivi", italic=True, size=7.5, color=MUTED_SOFT)
    lignes_plan = []
    for i, t in enumerate(ch.get("taches") or [], 1):
        duree = "JALON" if t.get("is_milestone") else f"{t.get('duree', 0)} j"
        if t.get("done"):
            etat = "Fait le " + str(t.get("done_date") or "")
        elif t.get("start_date"):
            etat = "En cours depuis le " + str(t.get("start_date"))
        else:
            etat = "À faire"
        lignes_plan.append((str(i), t.get("label") or "", duree, etat))
    _table_charte(doc, [("N°", 12, "center"), ("Lot / jalon", 104, "body"),
                        ("Durée", 22, "mono"), ("État", 48, "body")], lignes_plan)

    _eyebrow(doc, "Registre des risques", "Risk register")
    risques = ch.get("risques") or []
    if risques:
        st_lbl = {"ouvert": "Ouvert", "maitrise": "Maîtrisé",
                  "avere": "Avéré", "clos": "Clos"}
        lignes_rk = []
        for rk in risques:
            pgc = int(rk.get("probabilite") or 0) * int(rk.get("gravite") or 0)
            lignes_rk.append((rk.get("libelle") or "",
                              f"{rk.get('probabilite', '')}×{rk.get('gravite', '')} = {pgc}",
                              rk.get("parade") or "",
                              rk.get("responsable") or "",
                              st_lbl.get(rk.get("statut"), rk.get("statut") or "")))
        _table_charte(doc, [("Risque", 62, "body"), ("P×G", 20, "mono"),
                            ("Parade", 62, "body"), ("Resp.", 22, "body"),
                            ("État", 20, "body")], lignes_rk)
    else:
        p = doc.add_paragraph()
        _run(p, "Aucun risque enregistré au registre à la date d'édition.",
             italic=True, size=8.5, color=MUTED_SOFT)

    # ---- sommaire (titre centre, gras italique : le style du RFF) -----------
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    _run(p, "Table des matières", bold=True, italic=True, size=16, color=INK,
         caps=True, ls=16)
    _toc(doc)
    doc.add_page_break()

    # ---- sections -----------------------------------------------------------
    for i, s in enumerate(sections, 1):
        h = doc.add_heading(level=1)
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        _run(h, f"{i}. {s.get('titre') or 'Section'}", bold=True, size=13,
             color=INK, font=DIN, caps=True, ls=16)
        _p_border(h, sz=6, color="444444", space=3)
        for ligne in (s.get("corps") or "").split("\n"):
            _para_corps(doc, ligne)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    _p_border(p, edge="top", sz=4, color=HEX_HAIR, space=4)
    _run(p, "Document éditable : retouchez-le dans Word puis réimportez-le dans le "
            "suivi des chantiers. Ne supprimez pas les titres de section, ils servent "
            "à recoller le texte au bon endroit.", italic=True, size=7.5,
         color=MUTED_SOFT)

    buf = io.BytesIO()
    doc.save(buf)
    data = _injecte_props(buf.getvalue(), ch.get("id", ""),
                          [s.get("id", "") for s in sections])
    return data, _nom_fichier(ch, cdc)


# --------------------------------------------------------------------------- #
# Relecture
# --------------------------------------------------------------------------- #
_RE_P = re.compile(r"<w:p\b[^>]*>(.*?)</w:p>|<w:p\b[^>]*/>", re.S)
_RE_STYLE = re.compile(r'<w:pStyle\s+w:val="([^"]+)"')
_RE_T = re.compile(r"<w:t\b[^>]*>(.*?)</w:t>", re.S)
_RE_BR = re.compile(r"<w:br\b[^>]*/?>")
_RE_TAG = re.compile(r"<[^>]+>")
_RE_PROP = re.compile(r'<property[^>]*name="([^"]+)"[^>]*>\s*<vt:lpwstr>(.*?)</vt:lpwstr>', re.S)
_RE_NUM = re.compile(r"^\s*\d+(?:\.\d+)*[.)]?\s+")


def _deentite(s: str) -> str:
    s = _RE_TAG.sub("", s)
    for a, b in (("&lt;", "<"), ("&gt;", ">"), ("&quot;", '"'), ("&apos;", "'"), ("&amp;", "&")):
        s = s.replace(a, b)
    return s


def _detexte(xml_p: str) -> str:
    xml_p = _RE_BR.sub("\n", xml_p)
    return "".join(_deentite(m.group(1)) for m in _RE_T.finditer(xml_p))


def parse(data: bytes) -> dict:
    """Relit un .docx produit par build(). Leve ValueError si inexploitable."""
    try:
        z = zipfile.ZipFile(io.BytesIO(data))
        document = z.read("word/document.xml").decode("utf-8", "replace")
    except KeyError:
        raise ValueError("Ce fichier n'est pas un document Word (word/document.xml absent).")
    except zipfile.BadZipFile:
        raise ValueError("Fichier illisible : ce n'est pas un .docx. "
                         "Enregistre bien au format Word (.docx), pas .doc ni .pdf.")

    chantier_id, ids = None, []
    try:
        custom = z.read("docProps/custom.xml").decode("utf-8", "replace")
        props = {m.group(1): m.group(2) for m in _RE_PROP.finditer(custom)}
        chantier_id = props.get(PROP_CHANTIER) or None
        ids = [x for x in (props.get(PROP_SECTIONS) or "").split(",") if x]
    except KeyError:
        pass                    # document reenregistre autrement : on retombe sur les titres

    # Tout ce qui precede le premier titre de section (garde, sommaire, tables)
    # est ignore : seule la matiere des sections fait l'aller-retour.
    sections, courante = [], None
    for m in _RE_P.finditer(document):
        bloc = m.group(1) or ""
        style = _RE_STYLE.search(bloc)
        texte = _detexte(bloc)
        if style and style.group(1) == STYLE_TITRE:
            titre = _RE_NUM.sub("", texte).strip()      # « 3. Périmètre » -> « Périmètre »
            courante = {"titre": titre, "lignes": []}
            sections.append(courante)
        elif courante is not None:
            courante["lignes"].append(texte)

    if not sections:
        raise ValueError("Aucune section trouvée. Les titres de section doivent rester "
                         "en style « Titre 1 » pour être reconnus.")

    sorties = []
    for i, sec in enumerate(sections):
        lignes = sec["lignes"]
        while lignes and not lignes[-1].strip():        # note de pied ajoutee par build()
            lignes.pop()
        if lignes and lignes[-1].startswith("Document éditable"):
            lignes.pop()
        corps = "\n".join(lignes).strip("\n")
        sorties.append({"id": ids[i] if i < len(ids) else None,
                        "titre": sec["titre"], "corps": corps})
    return {"chantier_id": chantier_id, "sections": sorties}
